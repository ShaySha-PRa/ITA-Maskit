"""Word (.docx) 读写。

python-docx 遍历段落 + 表格，每段文本走 mask_text_pii，原地写回保留格式。
"""
from __future__ import annotations

from pathlib import Path

from maskit.rules.defs import RuleSet
from maskit.text import mask_text_pii

try:
    import docx  # python-docx
except ImportError:  # pragma: no cover
    docx = None


def _mask_runs_paragraph(
    paragraph, ruleset: RuleSet, pepper: str | None, strategy: str,
    scan_names: bool = False, person_list: set[str] | None = None,
) -> None:
    """脱敏一个段落（保留段落级格式，文本按 runs 重写）。

    简化方案：拼接全部 run 文本 → 整体脱敏 → 写回第一个 run，其余 run 清空。
    会丢失 run 级格式（如加粗/斜体），但保留段落级格式，可接受。
    """
    if not paragraph.runs:
        return
    original = "".join(run.text for run in paragraph.runs)
    if not original.strip():
        return
    masked = mask_text_pii(original, ruleset, pepper, strategy, scan_names, person_list)
    if masked == original:
        return
    paragraph.runs[0].text = masked
    for run in paragraph.runs[1:]:
        run.text = ""


def _mask_table(
    table, ruleset: RuleSet, pepper: str | None, strategy: str,
    scan_names: bool = False, person_list: set[str] | None = None,
) -> None:
    """脱敏表格每个单元格。"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _mask_runs_paragraph(paragraph, ruleset, pepper, strategy, scan_names, person_list)


def mask_docx_file(
    input_path: str | Path,
    output_path: str | Path,
    ruleset: RuleSet,
    pepper: str | None,
    strategy: str = "mask",
    scan_names: bool = False,
    person_list: set[str] | None = None,
) -> int:
    """脱敏 .docx → .docx，返回段落数。"""
    src = Path(input_path)
    dst = Path(output_path)
    if not src.exists():
        raise FileNotFoundError(f"输入文件不存在: {src}")

    if docx is None:
        raise ValueError("需要安装 python-docx 才能处理 Word")

    try:
        document = docx.Document(str(src))
    except Exception as exc:
        raise ValueError(f"无法读取 Word 文件: {src} ({exc})") from exc

    count = 0
    # 正文段落
    for paragraph in document.paragraphs:
        _mask_runs_paragraph(paragraph, ruleset, pepper, strategy, scan_names, person_list)
        count += 1
    # 表格
    for table in document.tables:
        _mask_table(table, ruleset, pepper, strategy, scan_names, person_list)
        count += sum(len(r.cells) for r in table.rows)

    document.save(str(dst))
    return count
