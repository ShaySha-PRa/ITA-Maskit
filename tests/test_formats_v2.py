"""v2 格式测试：Excel/JSON/邮件/PDF/Word 端到端 + 统一入口分发。"""

import polars as pl
import pytest

from maskit.io import SUPPORTED_FORMATS, is_text_format, mask_file
from maskit.rules.loader import load_ruleset

# --- 文本引擎单测 ---


def test_mask_text_pii_mask():
    """文本扫描：mask 策略替换有精确正则的 PII，保留普通文字。"""
    from maskit.text import mask_text_pii

    rs = load_ruleset()
    text = "联系 alice@corp.example 或 138-0000-0000，服务器 10.1.2.3，张伟审批。"
    out = mask_text_pii(text, rs, None, "mask")
    assert "alice@corp.example" not in out
    assert "a***@corp.example" in out
    assert "138****0000" in out
    assert "*.*.*.*" in out
    assert "张伟" in out  # name 不参与全文扫描


def test_mask_text_pii_pseudo_deterministic():
    """文本扫描：pseudo 确定性。"""
    from maskit.text import mask_text_pii

    rs = load_ruleset()
    text = "联系 alice@corp.example 或 13800000000"
    a = mask_text_pii(text, rs, "secret", "pseudo")
    b = mask_text_pii(text, rs, "secret", "pseudo")
    assert a == b
    assert "alice@corp.example" not in a


def test_mask_text_pii_no_pepper_pseudo():
    """文本 pseudo 无 pepper → 报错。"""
    from maskit.text import mask_text_pii

    rs = load_ruleset()
    with pytest.raises(ValueError):
        mask_text_pii("alice@corp.example", rs, None, "pseudo")


# --- Excel ---


def test_excel_end_to_end(tmp_path):
    src = tmp_path / "in.xlsx"
    out = tmp_path / "out.xlsx"
    pl.DataFrame(
        {"name": ["张伟", "李娜"], "email": ["a@b.com", "c@d.com"], "ip": ["10.1.2.3", "10.1.2.4"]}
    ).write_excel(src)
    rows = mask_file(src, out, load_ruleset(), None)
    assert rows == 2
    masked = pl.read_excel(out)
    assert masked["name"].to_list() == ["张*", "李*"]
    assert masked["ip"].to_list() == ["*.*.*.*", "*.*.*.*"]


def test_excel_multi_sheet(tmp_path):
    """多 sheet Excel：每个 sheet 都脱敏，保留 sheet 结构。"""
    from openpyxl import Workbook, load_workbook

    src = tmp_path / "in.xlsx"
    out = tmp_path / "out.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "用户"
    ws1.append(["name", "email", "ip"])
    ws1.append(["张伟", "a@b.com", "10.1.2.3"])
    ws2 = wb.create_sheet("权限")
    ws2.append(["employee_id", "phone"])
    ws2.append(["EID-1001", "13800000000"])
    wb.save(str(src))

    rows = mask_file(src, out, load_ruleset(), None)
    assert rows == 2  # 两个 sheet 各 1 行

    wb2 = load_workbook(str(out))
    assert wb2.sheetnames == ["用户", "权限"]
    data1 = [[str(c.value) for c in row] for row in wb2["用户"].iter_rows()]
    data2 = [[str(c.value) for c in row] for row in wb2["权限"].iter_rows()]
    assert data1[1] == ["张*", "a***@b.com", "*.*.*.*"]
    assert data2[1] == ["EID-***", "138****0000"]


# --- JSON / JSONL ---


def test_jsonl_end_to_end(tmp_path):
    src = tmp_path / "in.ndjson"
    out = tmp_path / "out.ndjson"
    pl.DataFrame({"name": ["张伟"], "email": ["a@b.com"]}).write_ndjson(src)
    rows = mask_file(src, out, load_ruleset(), None)
    assert rows == 1
    masked = pl.read_ndjson(out)
    assert masked["name"].to_list() == ["张*"]


def test_json_array_end_to_end(tmp_path):
    src = tmp_path / "in.json"
    out = tmp_path / "out.json"
    import json

    src.write_text(json.dumps([{"name": "张伟", "email": "a@b.com"}]), encoding="utf-8")
    rows = mask_file(src, out, load_ruleset(), None)
    assert rows == 1
    # 输出为 NDJSON
    data = [json.loads(l) for l in out.read_text(encoding="utf-8").strip().splitlines()]
    assert data[0]["name"] == "张*"


# --- 邮件 ---


def test_email_end_to_end(tmp_path):
    src = tmp_path / "in.eml"
    out = tmp_path / "out.eml"
    src.write_text(
        """From: 张伟 <zhangwei@corp.example>
To: auditor@corp.example
Subject: 审批
Date: Mon, 05 Aug 2026 10:00:00 +0800

请审批 alice@corp.example 权限，电话 138-0000-0000，服务器 10.1.2.3。
""",
        encoding="utf-8",
    )
    rows = mask_file(src, out, load_ruleset(), None, strategy="mask")
    assert rows == 1
    import email

    msg = email.message_from_string(out.read_text(encoding="utf-8"))
    assert "alice@corp.example" not in msg.get_payload(decode=True).decode("utf-8", "ignore")
    assert "zhangwei@corp.example" not in str(msg["From"])  # From 地址被脱敏


def test_email_deterministic(tmp_path):
    """邮件 pseudo 确定性：同输入同 pepper → 逐字节一致。"""
    src = tmp_path / "in.eml"
    src.write_text(
        "From: a@b.com\nTo: c@d.com\nSubject: x\n\nalice@corp.example 13800000000\n",
        encoding="utf-8",
    )
    out1 = tmp_path / "o1.eml"
    out2 = tmp_path / "o2.eml"
    mask_file(src, out1, load_ruleset(), "pep", strategy="pseudo")
    mask_file(src, out2, load_ruleset(), "pep", strategy="pseudo")
    assert out1.read_bytes() == out2.read_bytes()


# --- Word ---


def test_docx_end_to_end(tmp_path):
    from docx import Document

    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("申请人 alice@corp.example 电话 138-0000-0000")
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "10.1.2.3"
    doc.save(str(src))
    mask_file(src, out, load_ruleset(), None, strategy="mask")
    r = Document(str(out))
    assert "alice@corp.example" not in r.paragraphs[0].text
    assert "*.*.*.*" in r.tables[0].rows[0].cells[0].text


# --- PDF ---


def test_pdf_end_to_end(tmp_path):
    from pypdf import PdfReader
    from reportlab.pdfgen import canvas

    src = tmp_path / "in.pdf"
    out = tmp_path / "out.pdf"
    c = canvas.Canvas(str(src))
    c.drawString(100, 750, "Contact alice@corp.example phone 138-0000-0000")
    c.save()
    mask_file(src, out, load_ruleset(), None, strategy="mask")
    text = PdfReader(str(out)).pages[0].extract_text()
    assert "alice@corp.example" not in text
    assert "138" in text


# --- 统一入口分发 ---


def test_mask_file_unsupported_format(tmp_path):
    src = tmp_path / "in.xyz"
    src.write_text("data", encoding="utf-8")
    with pytest.raises(ValueError, match="不支持的输入格式"):
        mask_file(src, tmp_path / "o.xyz", load_ruleset(), None)


def test_supported_formats():
    assert ".csv" in SUPPORTED_FORMATS
    assert ".xlsx" in SUPPORTED_FORMATS
    assert ".eml" in SUPPORTED_FORMATS
    assert ".pdf" in SUPPORTED_FORMATS
    assert ".docx" in SUPPORTED_FORMATS


def test_is_text_format():
    assert is_text_format("x.eml")
    assert is_text_format("x.pdf")
    assert is_text_format("x.docx")
    assert not is_text_format("x.csv")
    assert not is_text_format("x.xlsx")
